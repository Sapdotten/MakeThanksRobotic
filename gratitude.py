import docx
from docx.shared import Pt


def gratitude():
    doc = docx.Document("SampleThanks.docx")  # макет документа в папке с кодом

    # стили шрифта
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(14)

    # Обучающимся института
    institute = doc.paragraphs[3]
    run = institute.add_run(input("Введите институт: ") + ":")
    run.bold = True

    # Заполнение таблицы
    fio = doc.tables[2]

    while True:
        num_rows = len(fio.rows)
        cells = fio.add_row().cells
        cells[0].text = str(num_rows) + "."  # порядковый номер
        cells[1].text = input("Введите ФИО: ")  # ФИО
        cells[2].text = input("Введите Группу: ")  # Группа
        next_step = input("Для продолжения нажмите Enter, Чтобы закончить напишите 'e', : ")
        if next_step.lower() == "e" or next_step.lower() == "е":  # e на русском и на англ
            break

    # Организация и проведение
    organization = input("за организацию и проведение: ")
    organization_print = doc.paragraphs[6]
    organization_print.add_run(organization + ".")

    doc.save(
        f"ReadyThanks\\Благотворительное письмо {organization}.docx")  # необходимо создать папку "ReadyThanks" в корне проекта


gratitude()
